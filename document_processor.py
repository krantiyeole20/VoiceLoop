"""
Document Processing Module for RAG Implementation

Extracts text from PDF, DOCX, MD, TXT files and caches for interview context injection.
Privacy-first: Only stores extracted text, never raw files.
"""

import os
import hashlib
import logging
import time
from typing import Optional, Dict, List, Union, BinaryIO
from dataclasses import dataclass, field
from pathlib import Path
import io

logger = logging.getLogger(__name__)


@dataclass
class DocumentMetadata:
    """Metadata for cached documents."""
    filename: str
    document_type: str  # 'resume', 'job_description', 'portfolio'
    uploaded_at: float
    file_size: int
    extraction_method: str = ""
    char_count: int = 0


class DocumentProcessor:
    """
    Handles document text extraction and caching for RAG.
    
    Supports: PDF, DOCX, MD, TXT
    Privacy-first: Only caches extracted text, not original files.
    """

    def __init__(self):
        """Initialize document processor with in-memory cache."""
        self.cache: Dict[str, dict] = {}
        logger.info("[DOC_PROCESSOR] Document processor initialized")

    def extract_text(self, file_or_path: Union[str, Path, BinaryIO], filename: str = None) -> str:
        """
        Extract text from a file based on its extension.
        
        Args:
            file_or_path: File path string, Path object, or file-like object
            filename: Optional filename (required if file_or_path is file-like object)
            
        Returns:
            Extracted and cleaned text
        """
        try:
            # Determine file extension
            if isinstance(file_or_path, (str, Path)):
                filepath = Path(file_or_path)
                filename = filepath.name
                ext = filepath.suffix.lower()
                
                if not filepath.exists():
                    logger.error(f"[DOC_PROCESSOR] File not found: {filepath}")
                    return ""
                    
                with open(filepath, 'rb') as f:
                    content = f.read()
            else:
                # File-like object
                if not filename:
                    logger.error("[DOC_PROCESSOR] Filename required for file-like objects")
                    return ""
                ext = Path(filename).suffix.lower()
                content = file_or_path.read()
                if hasattr(file_or_path, 'seek'):
                    file_or_path.seek(0)
            
            logger.info(f"[DOC_PROCESSOR] Extracting text from: {filename} (type: {ext})")
            
            # Route to appropriate extractor
            if ext == '.pdf':
                text = self._extract_pdf(content)
            elif ext == '.docx':
                text = self._extract_docx(content)
            elif ext == '.doc':
                text = self._extract_doc_fallback(content, filename)
            elif ext in ['.md', '.txt', '.text']:
                text = self._extract_plaintext(content)
            else:
                logger.warning(f"[DOC_PROCESSOR] Unsupported file type: {ext}")
                return f"[Unsupported file type: {ext}. Supported types: PDF, DOCX, MD, TXT]"
            
            cleaned = self.clean_text(text)
            logger.info(f"[DOC_PROCESSOR] Extracted {len(cleaned)} characters from {filename}")
            return cleaned
            
        except Exception as e:
            logger.error(f"[DOC_PROCESSOR] Extraction error: {e}", exc_info=True)
            return f"[Error extracting text: {str(e)}]"

    def _extract_pdf(self, content: bytes) -> str:
        """Extract text from PDF using PyPDF2."""
        try:
            import PyPDF2
            
            pdf_file = io.BytesIO(content)
            reader = PyPDF2.PdfReader(pdf_file)
            
            text_parts = []
            for page_num, page in enumerate(reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
                except Exception as e:
                    logger.warning(f"[DOC_PROCESSOR] Error extracting page {page_num}: {e}")
                    continue
            
            text = "\n\n".join(text_parts)
            logger.debug(f"[DOC_PROCESSOR] PDF extracted {len(reader.pages)} pages")
            return text
            
        except ImportError:
            logger.error("[DOC_PROCESSOR] PyPDF2 not installed. Run: pip install PyPDF2")
            return "[PDF extraction requires PyPDF2. Please install it.]"
        except Exception as e:
            logger.error(f"[DOC_PROCESSOR] PDF extraction error: {e}", exc_info=True)
            return f"[PDF extraction failed: {str(e)}]"

    def _extract_docx(self, content: bytes) -> str:
        """Extract text from DOCX using python-docx."""
        try:
            from docx import Document
            
            docx_file = io.BytesIO(content)
            doc = Document(docx_file)
            
            text_parts = []
            
            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(" | ".join(row_text))
            
            text = "\n".join(text_parts)
            logger.debug(f"[DOC_PROCESSOR] DOCX extracted {len(doc.paragraphs)} paragraphs")
            return text
            
        except ImportError:
            logger.error("[DOC_PROCESSOR] python-docx not installed. Run: pip install python-docx")
            return "[DOCX extraction requires python-docx. Please install it.]"
        except Exception as e:
            logger.error(f"[DOC_PROCESSOR] DOCX extraction error: {e}", exc_info=True)
            return f"[DOCX extraction failed: {str(e)}]"

    def _extract_doc_fallback(self, content: bytes, filename: str) -> str:
        """Best-effort extraction for legacy .doc files."""
        logger.warning(f"[DOC_PROCESSOR] Legacy .doc format detected: {filename}")
        
        # Try to extract any readable text from binary
        try:
            # Simple approach: decode as latin-1 and look for text patterns
            text = content.decode('latin-1', errors='ignore')
            
            # Filter to printable ASCII-ish characters
            printable_chars = []
            word_buffer = []
            
            for char in text:
                if char.isprintable() or char in '\n\t':
                    word_buffer.append(char)
                else:
                    if len(word_buffer) >= 3:  # Only keep sequences of 3+ chars
                        printable_chars.extend(word_buffer)
                    word_buffer = []
            
            if word_buffer:
                printable_chars.extend(word_buffer)
            
            extracted = ''.join(printable_chars)
            
            if len(extracted) > 100:
                return f"[Legacy .doc format - partial extraction]\n{extracted}"
            else:
                return (
                    f"[Legacy .doc format not fully supported. "
                    f"Please convert '{filename}' to .docx or .pdf for better results.]"
                )
                
        except Exception as e:
            logger.error(f"[DOC_PROCESSOR] .doc fallback error: {e}")
            return f"[Legacy .doc format not supported. Please convert to .docx or .pdf]"

    def _extract_plaintext(self, content: bytes) -> str:
        """Extract text from plaintext files (MD, TXT)."""
        # Try common encodings
        encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
        
        for encoding in encodings:
            try:
                return content.decode(encoding)
            except UnicodeDecodeError:
                continue
        
        # Fallback to utf-8 with replacement
        return content.decode('utf-8', errors='replace')

    def clean_text(self, text: str) -> str:
        """
        Clean extracted text for better processing.

        Operations:
        - Normalize whitespace
        - Remove extra line breaks
        - Fix common PDF extraction artifacts
        """
        if not text:
            return ""

        # Replace common artifacts
        text = text.replace('\x00', '')  # Null bytes
        text = text.replace('\r\n', '\n')  # Normalize line endings
        text = text.replace('\r', '\n')
        
        # Fix hyphenation at line breaks (common in PDFs)
        import re
        text = re.sub(r'(\w)-\n(\w)', r'\1\2', text)
        
        # Normalize multiple newlines to max 2
        text = re.sub(r'\n{3,}', '\n\n', text)
        
        # Normalize multiple spaces
        text = re.sub(r'[ \t]+', ' ', text)
        
        # Clean up lines
        lines = []
        for line in text.split('\n'):
            stripped = line.strip()
            if stripped:
                lines.append(stripped)
            elif lines and lines[-1] != '':
                lines.append('')
        
        text = '\n'.join(lines).strip()

        logger.debug(f"[DOC_PROCESSOR] Cleaned text: {len(text)} characters")
        return text

    def cache_document(
        self,
        text: str,
        metadata: DocumentMetadata
    ) -> str:
        """
        Cache document with metadata for quick retrieval.
        Uses MD5 hash as cache key for deduplication.

        Args:
            text: Document text
            metadata: Document metadata

        Returns:
            Cache key (MD5 hash)
        """
        if not text:
            logger.warning("[DOC_PROCESSOR] Attempted to cache empty document")
            return ""

        # Generate cache key
        key = hashlib.md5(text.encode()).hexdigest()

        # Check if already cached
        if key in self.cache:
            logger.info(f"[DOC_PROCESSOR] Document already cached: {key}")
            return key

        # Update metadata with char count
        metadata.char_count = len(text)

        # Store in cache
        self.cache[key] = {
            'text': text,
            'metadata': metadata,
            'text_length': len(text),
        }

        logger.info(
            f"[DOC_PROCESSOR] Cached document: {key} "
            f"({metadata.filename}, {len(text)} chars, type: {metadata.document_type})"
        )

        return key

    def get_cached_document(self, cache_key: str) -> Optional[dict]:
        """
        Retrieve a cached document by key.

        Args:
            cache_key: The cache key from cache_document

        Returns:
            Cached document dict or None
        """
        return self.cache.get(cache_key)

    def get_cached_text(self, cache_key: str) -> str:
        """
        Retrieve just the text from a cached document.

        Args:
            cache_key: The cache key from cache_document

        Returns:
            Cached text or empty string
        """
        doc = self.cache.get(cache_key)
        if doc:
            return doc.get('text', '')
        return ''

    def retrieve_relevant_context(
        self,
        query: str,
        cached_key: Optional[str] = None,
        max_length: int = 500
    ) -> str:
        """
        Retrieve relevant context for a query.

        Current: Returns first N characters (simple implementation)
        Future: Will implement semantic search using cosine similarity

        Args:
            query: Query text to find relevant context for
            cached_key: Cache key of document to search in
            max_length: Maximum length of returned context

        Returns:
            Relevant text snippet
        """
        if not cached_key or cached_key not in self.cache:
            logger.warning(
                f"[DOC_PROCESSOR] Invalid cache key or document not found: {cached_key}"
            )
            return ""

        doc = self.cache[cached_key]
        text = doc['text']

        # Simple implementation: Return first N characters
        # TODO: Implement semantic search with embeddings
        context = text[:max_length]
        if len(text) > max_length:
            context += "..."

        logger.debug(
            f"[DOC_PROCESSOR] Retrieved context: {len(context)} chars "
            f"(query: '{query[:50]}...')"
        )

        return context

    def get_cache_stats(self) -> dict:
        """Get statistics about cached documents."""
        total_docs = len(self.cache)
        total_chars = sum(doc['text_length'] for doc in self.cache.values())
        
        by_type = {}
        for doc in self.cache.values():
            doc_type = doc.get('metadata', {})
            if isinstance(doc_type, DocumentMetadata):
                doc_type = doc_type.document_type
            else:
                doc_type = doc_type.get('document_type', 'unknown')
            by_type[doc_type] = by_type.get(doc_type, 0) + 1

        return {
            'total_documents': total_docs,
            'total_characters': total_chars,
            'by_type': by_type,
            'cache_keys': list(self.cache.keys())
        }

    def clear_cache(self):
        """Clear the document cache."""
        count = len(self.cache)
        self.cache.clear()
        logger.info(f"[DOC_PROCESSOR] Cache cleared ({count} documents removed)")

    def remove_cached(self, cache_key: str) -> bool:
        """Remove a specific document from cache."""
        if cache_key in self.cache:
            del self.cache[cache_key]
            logger.info(f"[DOC_PROCESSOR] Removed cached document: {cache_key}")
            return True
        return False


# Global instance for easy access
doc_processor = DocumentProcessor()